#!/usr/bin/env python3
"""
Report exporter for converting markdown research reports to PDF and Word formats.

Usage:
    from core.report_exporter import ReportExporter

    exporter = ReportExporter()
    exporter.to_pdf("report.md", "report.pdf")
    exporter.to_word("report.md", "report.docx")

    # Or export both at once
    exporter.export_all("report.md", output_dir="./exports")
"""

import logging
import re
from pathlib import Path
from typing import Optional, Union
from datetime import datetime

logger = logging.getLogger(__name__)

# CSS for PDF styling - professional report appearance
PDF_STYLESHEET = """
@page {
    size: letter;
    margin: 1in;
    @bottom-center {
        content: "Page " counter(page) " of " counter(pages);
        font-size: 9pt;
        color: #666;
    }
}

body {
    font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.5;
    color: #333;
    max-width: 100%;
}

h1 {
    font-size: 24pt;
    color: #1a1a2e;
    border-bottom: 2px solid #1a1a2e;
    padding-bottom: 10px;
    margin-top: 0;
}

h2 {
    font-size: 18pt;
    color: #16213e;
    border-bottom: 1px solid #ddd;
    padding-bottom: 5px;
    margin-top: 25px;
}

h3 {
    font-size: 14pt;
    color: #0f3460;
    margin-top: 20px;
}

h4 {
    font-size: 12pt;
    color: #333;
    font-style: italic;
}

p {
    margin: 10px 0;
    text-align: justify;
}

ul, ol {
    margin: 10px 0 10px 20px;
}

li {
    margin: 5px 0;
}

a {
    color: #0066cc;
    text-decoration: none;
}

a:hover {
    text-decoration: underline;
}

blockquote {
    border-left: 4px solid #0066cc;
    margin: 15px 0;
    padding: 10px 20px;
    background: #f8f9fa;
    font-style: italic;
}

code {
    background: #f4f4f4;
    padding: 2px 6px;
    border-radius: 3px;
    font-family: 'Consolas', 'Monaco', monospace;
    font-size: 10pt;
}

pre {
    background: #f4f4f4;
    padding: 15px;
    border-radius: 5px;
    overflow-x: auto;
    font-size: 9pt;
}

table {
    border-collapse: collapse;
    width: 100%;
    margin: 15px 0;
}

th, td {
    border: 1px solid #ddd;
    padding: 8px 12px;
    text-align: left;
}

th {
    background: #f8f9fa;
    font-weight: bold;
}

tr:nth-child(even) {
    background: #fafafa;
}

hr {
    border: none;
    border-top: 1px solid #ddd;
    margin: 25px 0;
}

/* Entity network styling */
strong {
    color: #1a1a2e;
}

/* Source citations */
em {
    color: #666;
}
"""


class ReportExporter:
    """
    Export markdown research reports to PDF and Word formats.

    Features:
    - Professional styling for PDF output
    - Preserves links and formatting in Word
    - Handles research report structure (headers, citations, entity networks)
    """

    def __init__(self):
        """Initialize the exporter."""
        self._markdown = None
        self._weasyprint = None
        self._docx = None

    def _ensure_markdown(self):
        """Lazy import markdown library."""
        if self._markdown is None:
            import markdown
            self._markdown = markdown.Markdown(
                extensions=['tables', 'fenced_code', 'toc']
            )
        return self._markdown

    def _ensure_weasyprint(self):
        """Lazy import weasyprint library."""
        if self._weasyprint is None:
            from weasyprint import HTML, CSS
            self._weasyprint = (HTML, CSS)
        return self._weasyprint

    def _ensure_docx(self):
        """Lazy import python-docx library."""
        if self._docx is None:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            self._docx = {
                'Document': Document,
                'Inches': Inches,
                'Pt': Pt,
                'WD_ALIGN_PARAGRAPH': WD_ALIGN_PARAGRAPH,
                'WD_STYLE_TYPE': WD_STYLE_TYPE
            }
        return self._docx

    def _read_markdown(self, source: Union[str, Path]) -> str:
        """
        Read markdown content from file or return string directly.

        Args:
            source: File path or markdown string

        Returns:
            Markdown content as string
        """
        source_path = Path(source)
        if source_path.exists() and source_path.is_file():
            with open(source_path, 'r', encoding='utf-8') as f:
                return f.read()
        return str(source)

    def _markdown_to_html(self, markdown_content: str) -> str:
        """
        Convert markdown to HTML.

        Args:
            markdown_content: Markdown text

        Returns:
            HTML string
        """
        md = self._ensure_markdown()
        md.reset()
        html_body = md.convert(markdown_content)

        # Wrap in full HTML document
        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Research Report</title>
</head>
<body>
{html_body}
</body>
</html>"""
        return html

    def to_pdf(
        self,
        source: Union[str, Path],
        output_path: Union[str, Path],
        title: Optional[str] = None
    ) -> Path:
        """
        Export markdown to PDF.

        Args:
            source: Markdown file path or content string
            output_path: Output PDF file path
            title: Optional title for the document

        Returns:
            Path to generated PDF file
        """
        HTML, CSS = self._ensure_weasyprint()

        markdown_content = self._read_markdown(source)
        html_content = self._markdown_to_html(markdown_content)

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Generate PDF with styling
        html_doc = HTML(string=html_content)
        css = CSS(string=PDF_STYLESHEET)

        html_doc.write_pdf(
            str(output_path),
            stylesheets=[css]
        )

        logger.info(f"PDF exported to: {output_path}")
        return output_path

    def to_word(
        self,
        source: Union[str, Path],
        output_path: Union[str, Path],
        title: Optional[str] = None
    ) -> Path:
        """
        Export markdown to Word document.

        Args:
            source: Markdown file path or content string
            output_path: Output Word file path
            title: Optional title for the document

        Returns:
            Path to generated Word file
        """
        docx_lib = self._ensure_docx()
        Document = docx_lib['Document']
        Pt = docx_lib['Pt']

        markdown_content = self._read_markdown(source)
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Parse markdown and add to document
        self._parse_markdown_to_docx(doc, markdown_content)

        doc.save(str(output_path))

        logger.info(f"Word document exported to: {output_path}")
        return output_path

    def _parse_markdown_to_docx(self, doc, markdown_content: str):
        """
        Parse markdown content and add to Word document.

        This is a simplified parser that handles common markdown elements.
        """
        docx_lib = self._ensure_docx()
        Pt = docx_lib['Pt']

        lines = markdown_content.split('\n')
        in_code_block = False
        code_block_content = []
        in_blockquote = False
        blockquote_content = []

        i = 0
        while i < len(lines):
            line = lines[i]

            # Code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End code block
                    code_text = '\n'.join(code_block_content)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Pt(20)
                    code_block_content = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue

            # Blockquotes
            if line.strip().startswith('>'):
                quote_text = line.strip()[1:].strip()
                if quote_text:
                    p = doc.add_paragraph(quote_text, style='Quote')
                i += 1
                continue

            # Headers
            if line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:].strip(), level=4)

            # Horizontal rules
            elif line.strip() in ['---', '***', '___']:
                p = doc.add_paragraph()
                p.add_run('_' * 50)

            # Unordered lists
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:].strip()
                text = self._clean_markdown_formatting(text)
                doc.add_paragraph(text, style='List Bullet')

            # Ordered lists
            elif re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                text = self._clean_markdown_formatting(text)
                doc.add_paragraph(text, style='List Number')

            # Regular paragraphs
            elif line.strip():
                text = self._clean_markdown_formatting(line)
                doc.add_paragraph(text)

            # Empty lines (paragraph breaks)
            else:
                pass  # Skip empty lines

            i += 1

    def _clean_markdown_formatting(self, text: str) -> str:
        """
        Clean markdown formatting from text for Word output.

        Converts markdown links to plain text with URLs, removes emphasis markers.
        """
        # Convert links [text](url) to "text (url)"
        text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', text)

        # Remove bold markers
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'__([^_]+)__', r'\1', text)

        # Remove italic markers
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        text = re.sub(r'_([^_]+)_', r'\1', text)

        # Remove inline code markers
        text = re.sub(r'`([^`]+)`', r'\1', text)

        return text

    def export_all(
        self,
        source: Union[str, Path],
        output_dir: Union[str, Path],
        base_name: Optional[str] = None,
        formats: Optional[list] = None
    ) -> dict:
        """
        Export markdown to multiple formats.

        Args:
            source: Markdown file path or content string
            output_dir: Directory to save exported files
            base_name: Base filename (without extension). If None, uses 'report'
            formats: List of formats to export. Default: ['pdf', 'docx']

        Returns:
            Dict mapping format to output path
        """
        if formats is None:
            formats = ['pdf', 'docx']

        if base_name is None:
            # Try to get base name from source if it's a file
            source_path = Path(source)
            if source_path.exists():
                base_name = source_path.stem
            else:
                base_name = 'report'

        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        results = {}

        if 'pdf' in formats:
            pdf_path = output_dir / f"{base_name}.pdf"
            results['pdf'] = self.to_pdf(source, pdf_path)

        if 'docx' in formats or 'word' in formats:
            docx_path = output_dir / f"{base_name}.docx"
            results['docx'] = self.to_word(source, docx_path)

        return results


def export_research_report(
    report_dir: Union[str, Path],
    formats: Optional[list] = None
) -> dict:
    """
    Convenience function to export a research report from its output directory.

    Args:
        report_dir: Path to research output directory containing report.md
        formats: List of formats to export. Default: ['pdf', 'docx']

    Returns:
        Dict mapping format to output path
    """
    report_dir = Path(report_dir)
    report_md = report_dir / "report.md"

    if not report_md.exists():
        raise FileNotFoundError(f"No report.md found in {report_dir}")

    exporter = ReportExporter()
    return exporter.export_all(
        source=report_md,
        output_dir=report_dir,
        base_name="report",
        formats=formats
    )


# CLI interface
if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Export markdown research reports to PDF and Word formats"
    )
    parser.add_argument(
        "source",
        help="Path to markdown file or research output directory"
    )
    parser.add_argument(
        "-o", "--output",
        help="Output directory (default: same as source)"
    )
    parser.add_argument(
        "-f", "--formats",
        nargs="+",
        default=["pdf", "docx"],
        choices=["pdf", "docx", "word"],
        help="Output formats (default: pdf docx)"
    )
    parser.add_argument(
        "-n", "--name",
        help="Base filename for output files (default: report)"
    )

    args = parser.parse_args()

    source = Path(args.source)

    # Determine if source is a directory or file
    if source.is_dir():
        report_md = source / "report.md"
        if not report_md.exists():
            print(f"Error: No report.md found in {source}")
            exit(1)
        source = report_md
        output_dir = args.output or source.parent
    else:
        output_dir = args.output or source.parent

    exporter = ReportExporter()
    results = exporter.export_all(
        source=source,
        output_dir=output_dir,
        base_name=args.name,
        formats=args.formats
    )

    print("\nExported reports:")
    for fmt, path in results.items():
        print(f"  {fmt.upper()}: {path}")
